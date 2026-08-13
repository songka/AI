from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "handover" / "AI-Coding-Session-Manager-任务交接书.md"
OUTPUT = ROOT / "docs" / "handover" / "AI-Coding-Session-Manager-任务交接书.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
DARK = "203040"
MID = "52606D"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_inches: list[float], indent_dxa: int = 100) -> None:
    """Write exact DXA widths so Word, LibreOffice, and geometry audits agree."""
    widths = [round(width * 1440) for width in widths_inches]
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_header(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def set_run_font(run, latin="Calibri", east_asia="Microsoft YaHei") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def add_inline(paragraph, text: str, *, base_size=10.5, color=DARK) -> None:
    # Render inline-code spans as lightly shaded monospace text.
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", "Microsoft YaHei")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLUE)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EAF1F6")
            run._r.get_or_add_rPr().append(shd)
        else:
            # Basic bold support, kept intentionally conservative.
            bold_parts = re.split(r"(\*\*[^*]+\*\*)", part)
            for bold_part in bold_parts:
                if not bold_part:
                    continue
                is_bold = bold_part.startswith("**") and bold_part.endswith("**")
                run = paragraph.add_run(bold_part[2:-2] if is_bold else bold_part)
                set_run_font(run)
                run.bold = is_bold
                run.font.size = Pt(base_size)
                run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.76)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in (
        ("Title", 25, BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 12.5, BLUE, 10, 4),
        ("Heading 3", 11, MID, 8, 3),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.12

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.74))
    table.autofit = False
    set_table_geometry(table, [4.7, 2.04], 90)
    set_repeat_table_header(table.rows[0])
    left, right = table.rows[0].cells
    left.text = "AI CODING SESSION MANAGER"
    right.text = "任务交接书  ·  V1.0"
    for idx, cell in enumerate((left, right)):
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell, 55, 90, 55, 90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            set_run_font(run)
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("内部项目交接资料   |   第 ")
    set_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    add_field(p, "PAGE")
    run = p.add_run(" 页")
    set_run_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)


def add_masthead(document: Document, lines: list[str]) -> None:
    title = lines[0].lstrip("# ").strip()
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(title)
    for run in p.runs:
        set_run_font(run)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)

    meta_rows = []
    for line in lines[1:]:
        clean = line.rstrip().rstrip("  ")
        if "：" in clean:
            label, value = clean.split("：", 1)
            meta_rows.append((label, value.strip("`")))
    table = document.add_table(rows=len(meta_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [1.15, 5.55])
    set_repeat_table_header(table.rows[0])
    for idx, (label, value) in enumerate(meta_rows):
        c0, c1 = table.rows[idx].cells
        c0.text = label
        c1.text = value
        set_cell_shading(c0, LIGHT_BLUE)
        set_cell_shading(c1, PALE_BLUE if idx % 2 == 0 else WHITE)
        for cell in (c0, c1):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor.from_string(DARK)
        c0.paragraphs[0].runs[0].bold = True
    document.add_paragraph()


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable = 6.72
    if col_count == 2:
        widths = [1.45, usable - 1.45]
    elif col_count == 3:
        widths = [1.35, 3.95, 1.42]
    else:
        widths = [usable / col_count] * col_count
    set_table_geometry(table, widths)
    for row_idx, values in enumerate(rows):
        row = table.rows[row_idx]
        prevent_row_split(row)
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            value = values[col_idx] if col_idx < len(values) else ""
            cell.text = ""
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, BLUE if row_idx == 0 else (LIGHT_GRAY if row_idx % 2 == 0 else WHITE))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, base_size=8.7, color=WHITE if row_idx == 0 else DARK)
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(document: Document, code_lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [6.7], 120)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "18212B")
    set_cell_margins(cell, 100, 120, 100, 120)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("\n".join(code_lines))
    set_run_font(run, "Consolas", "Microsoft YaHei")
    run.font.size = Pt(7.8)
    run.font.color.rgb = RGBColor.from_string("E8EEF4")
    prevent_row_split(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def create_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    # Title and metadata occupy the lines before the first level-two heading.
    first_section = next(i for i, line in enumerate(lines) if line.startswith("## "))
    add_masthead(document, [line for line in lines[:first_section] if line.strip()])

    i = first_section
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(document, rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            p = document.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(p, text, base_size={1: 16, 2: 12.5, 3: 11}[min(level, 3)], color=BLUE if level < 3 else MID)
            set_repeat_header(p)
            i += 1
            continue
        checkbox = re.match(r"^- \[([ xX])\]\s+(.+)$", stripped)
        if checkbox:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2.5)
            symbol = "☒" if checkbox.group(1).lower() == "x" else "☐"
            run = p.add_run(symbol + "  ")
            set_run_font(run, "Segoe UI Symbol", "Segoe UI Symbol")
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor.from_string(BLUE)
            add_inline(p, checkbox.group(2))
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = document.add_paragraph(style="List Number")
            add_inline(p, numbered.group(1))
            i += 1
            continue

        # Join adjacent plain-text lines into one paragraph.
        paragraph_lines = [stripped.rstrip("  ")]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate or candidate.startswith(("#", "|", "```", "- ")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate.rstrip("  "))
            i += 1
        p = document.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))

    document.core_properties.title = "AI Coding Session Manager 任务交接书"
    document.core_properties.subject = "项目状态、交付物、构建发布与后续任务交接"
    document.core_properties.author = "AI Coding Session Manager 项目组"
    document.core_properties.keywords = "WPF, .NET 8, Codex, Claude Code, OpenCode, 任务交接"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    create_document()
