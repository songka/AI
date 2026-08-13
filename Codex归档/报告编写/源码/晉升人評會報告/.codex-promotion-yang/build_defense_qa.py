from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\lfaf-test\Documents\报告编写\晉升人評會報告\楊敏銳_晉升課長_開放式答辯問答_多Agent模擬版.docx"

BLUE = "00457A"
DARK_BLUE = "1F4D78"
INK = "1F1F1F"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "FFC000"


def set_cell_text_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    ppr.append(OxmlElement("w:contextualSpacing"))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_repeat_font(run, size=9, color=MUTED)


def add_labeled_paragraph(doc, label, text, *, after=6, fill=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    label_run = p.add_run(label)
    set_repeat_font(label_run, bold=True, color=DARK_BLUE)
    text_run = p.add_run(text)
    set_repeat_font(text_run, color=INK, italic=italic)
    if fill:
        shade_paragraph(p, fill)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(7)
    return p


def add_question(doc, number, item):
    h = doc.add_paragraph(style="Heading 2")
    h.paragraph_format.keep_with_next = True
    run = h.add_run(f"Q{number}. {item['q']}")
    set_repeat_font(run, size=13, bold=True, color=BLUE)

    p = add_labeled_paragraph(doc, "回答主線｜", item["line"], after=5)
    p.paragraph_format.keep_with_next = True
    add_labeled_paragraph(doc, "示範回答｜", item["answer"], after=6)
    p = add_labeled_paragraph(doc, "可能追問｜", item["follow"], after=4)
    p.paragraph_format.keep_with_next = True
    add_labeled_paragraph(doc, "追問要點｜", item["follow_answer"], after=6)
    add_labeled_paragraph(
        doc,
        "證據邊界｜",
        item["risk"],
        after=9,
        fill=LIGHT_GRAY,
    )


sections = [
    (
        "一、角色躍遷與經營責任",
        [
            {
                "q": "你為什麼已經準備好從製造主管晉升課長？",
                "line": "用「既有證據 + 角色轉變 + 下一步責任」回答。",
                "answer": "我的基礎有三個：第一，具備品質與製造的完整經歷；第二，已帶領團隊完成六類重大專案；第三，已開始用六人調試售後團隊、標準化與安全庫存機制降低個人依賴。課長角色對我最大的改變，是從自己處理問題，轉為分解目標、配置資源、建立標準和培養梯隊，對團隊的持續交付負責。",
                "follow": "升任後如果你仍然每天救火，代表什麼？",
                "follow_answer": "代表授權、標準或人才梯隊仍不足。重大風險我會介入，但一般問題要交給責任人；每次救火後都要轉成流程、標準或培訓改善。",
                "risk": "不要只說「能力已經足夠」；要以PPT中的專案、團隊和機制作為證據。",
            },
            {
                "q": "主管與課長最大的差異是什麼？",
                "line": "從現場執行轉向目標、資源、風險與人才。",
                "answer": "主管更重視當天任務是否完成；課長還要看跨專案資源是否衝突、風險是否提前暴露、方法是否可複製、關鍵崗位是否有備援。我會把部門目標分到專案與責任人，用節點和指標追蹤，再把重複問題沉澱成標準與培訓內容。我的價值不再是自己解決多少問題，而是團隊離開我仍能穩定運作。",
                "follow": "你準備先授權哪一項工作？",
                "follow_answer": "選一項已有明確流程的工作，例如售後問題分級或月底庫存盤點；說清責任人、授權邊界、升級條件和檢核節點。",
                "risk": "PPT沒有具體人名，不要臨場虛構授權對象。",
            },
            {
                "q": "你目前距離一名稱職課長，最大的短板是什麼？",
                "line": "承認真實短板，並連到未來行動。",
                "answer": "我的優勢是品質與製造的完整視角，也有專案和現場管理經驗；短板是仍容易依靠自己處理問題，人才梯隊、授權機制與前端客戶需求的系統化掌握還要加強。因此下一階段我會把AB角、多能工、帶教，以及製造向前端延伸作為重點。我要證明的不是自己更忙，而是團隊更穩、更可複製。",
                "follow": "半年後用哪三個現象證明短板改善？",
                "follow_answer": "關鍵崗位有可獨立替補的B角；一般問題可由責任人閉環；製造能更早參與新專案風險評審。",
                "risk": "避免「溝通不夠」這類安全但空泛的回答，也不要把重大失誤包裝成短板。",
            },
            {
                "q": "上任後前90天，你最先做什麼？",
                "line": "30天盤點、60天定優先、90天驗證。",
                "answer": "前30天盤點專案、人員技能和主要交付風險；60天確定優先順序、責任人和檢核指標；90天完成第一輪AB角驗證、關鍵流程標準化及經營復盤。先從影響最大、單點風險最高的事項下手，不平均用力，再根據結果調整下一階段資源。",
                "follow": "團隊、技術、標準化只能先選一條，你選哪條？",
                "follow_answer": "原則上先抓團隊能力，因為技術與標準都要靠人落地；但若當期最大風險是重複異常或新製程判斷不足，優先級應隨經營瓶頸調整。",
                "risk": "這是未來計畫，不要表述成已完成。",
            },
        ],
    ),
    (
        "二、專案交付與資源配置",
        [
            {
                "q": "六類重大專案如何證明是你的管理成果，而不只是團隊把設備做完？",
                "line": "用代表專案講「風險—決策—協同—結果」。",
                "answer": "六類專案涵蓋衝壓、打磨、自動組裝等不同場景。我的作用不是替工程人員完成技術細節，而是從製造主管角度統籌人員、進度和現場問題，確保設備投入生產。答辯時我會選一個親自參與最深的專案，講清原始風險、我的判斷、協調了哪些資源，以及用計畫、問題關閉、驗收或投產資料證明結果。",
                "follow": "哪個專案最能證明你的貢獻？",
                "follow_answer": "只選最熟的一個；提前準備真實時間、關鍵異常、本人決策和驗收證據。",
                "risk": "PPT沒有交期達成率、預算、良率與客戶評價，不能臨場補造數字。",
            },
            {
                "q": "你給公司帶來的核心經營價值是什麼？",
                "line": "不要停在「完成了專案」，要連到交付、品質與複製。",
                "answer": "我的核心價值是把品質經驗帶到製造前端，在配線、裝配、調試與售後階段提前識別風險，同時把個人經驗轉成團隊和標準。現階段已形成六人調試售後團隊、三階段服務流程、配盤與標識標準、安全庫存清單。下一步會用準時交付、異常關閉、售後復發和人才替補等指標來衡量，而不是只報告完成了幾條線。",
                "follow": "目前哪一項價值最有證據？",
                "follow_answer": "選PPT有直接證據的團隊組建、24H響應機制、標準化或安全庫存，不要擴張到未披露的降本數字。",
                "risk": "不要宣稱節省多少成本或提升多少百分比，除非手上有正式資料。",
            },
            {
                "q": "為什麼要固定六人調試售後團隊？這不是增加編制嗎？",
                "line": "說清問題、分工、負荷與彈性使用。",
                "answer": "組建團隊是為了解決職責分散、售後依賴個人，以及電控人員頻繁被打斷的問題。六人承接調試與售後，統一調試期、保固期和保固外流程，對外建立穩定窗口，對內讓電控聚焦方案與程序。課長層面還要用案件量、工時、技能矩陣和一次解決情況檢驗配置；低峰期安排標準化、多能工和項目支援，避免人力閒置。",
                "follow": "如果公司要求減少一人，你怎麼安排？",
                "follow_answer": "先分析任務量與技能單點，保留關鍵技能和AB角，再用排班、遠端診斷與跨崗支援吸收缺口；安全與客戶停線不能作為減員代價。",
                "risk": "PPT沒有案件量和工時基線，只能說管理方法，不能證明六人是永久最優數量。",
            },
            {
                "q": "週末同時出現三起客戶停線，你如何排優先級？",
                "line": "安全、影響、時效、能力匹配四步判斷。",
                "answer": "先確認人身與設備安全，再按客戶停線範圍、產能影響、可否遠端止血和保固責任分級；接著按團隊技能與地點配置人員，必要時升級電控、機械或品質支援。每起問題都要有臨時措施、責任人與下一次更新時間，不能只看誰先打電話。",
                "follow": "最重要客戶與最嚴重安全風險衝突時怎麼選？",
                "follow_answer": "安全風險優先；同時向重要客戶透明說明資源與恢復方案，必要時啟動跨部門支援。",
                "risk": "不要承諾PPT未證明的值班制度或到場時限。",
            },
        ],
    ),
    (
        "三、製造、品質與流程閉環",
        [
            {
                "q": "24H響應是收到消息，還是已經解決問題？",
                "line": "明確定義：響應不等於結案。",
                "answer": "24H響應應定義為問題在24小時內被接收、初判、分級並給出處理安排，不等於所有問題24小時修復。真正要管理的是責任人、臨時止血、永久關閉和是否復發。後續我會分開追蹤響應時效、結案週期、重複問題和客戶回饋，避免團隊只做到回覆、沒有推進。",
                "follow": "原因暫時查不清時怎麼辦？",
                "follow_answer": "先保障安全與生產，給出臨時遏制方案；設定升級節點、責任人與下次更新時間，根因未關閉前持續跟蹤。",
                "risk": "PPT只有「24H響應」，沒有達成率或平均結案時間；不能說成24H解決。",
            },
            {
                "q": "標準化如何證明真正有效，而不只是看起來整齊？",
                "line": "有效標準要可執行、可檢查、可更新。",
                "answer": "我對標準化有效的定義，是不同人能按同一方法交付，並降低識別、學習和排查成本。目前已統一配盤布局、線色、號碼管標貼及關鍵走線與快換。下一步要把標準、檢查表、異常回饋和版本更新連起來，用現場稽核、重複異常、新人學習與維修排查資料驗證效果。",
                "follow": "資深員工認為標準拖慢效率，怎麼處理？",
                "follow_answer": "安全與法規不讓步；其他標準用現場數據判斷。若標準本身不合理就修訂，若只是習慣問題則透過培訓、稽核和責任要求落地。",
                "risk": "PPT沒有工時、返工率、稽核率與前後對比，不能聲稱已降低多少。",
            },
            {
                "q": "標準化與首台設備交期衝突時，你堅持還是讓步？",
                "line": "底線不讓步，一般標準分階段。",
                "answer": "涉及安全、法規和重大品質的標準不能讓步；一般性外觀或優化標準可分階段導入，但要留下經批准的偏差記錄、臨時控制與回歸時間。項目結束後必須復盤：是標準不合理、資源不足，還是執行問題，再更新標準，不能把一次例外變成長期習慣。",
                "follow": "誰有權批准偏差？",
                "follow_answer": "依公司流程由相應責任主管與品質/工程共同確認；自己不能單方面越權放行。",
                "risk": "不要宣稱已有正式偏差流程；可用「將依公司流程」守住邊界。",
            },
            {
                "q": "安全庫存降低缺料，也占用資金和空間，你如何平衡？",
                "line": "按重要性、交期、消耗、替代性和停線影響分類。",
                "answer": "安全庫存不是越多越好。現有改善是建立清單、分區分格、月底盤點，低於安全量開PR。下一步要按消耗頻率、採購週期、共用性、替代性和停線影響設定與調整安全量；高頻通用件可備安全量，低頻高價專用件則考慮供應商快速響應、寄售或分批交付，同時檢查呆滯。",
                "follow": "財務要求降庫存、製造要求加庫存，怎麼決策？",
                "follow_answer": "優先安全與關鍵路徑物料，用缺料影響、交期和替代方案共同評估，將資金占用與停線風險透明呈現後決策。",
                "risk": "PPT沒有庫存金額、周轉率和缺料降幅，不能宣稱已降本。",
            },
            {
                "q": "趕交期與守品質衝突時，你如何決策？",
                "line": "安全和重大品質紅線不帶病交付。",
                "answer": "我有品質與製造兩段經歷，不會把品質放到交付之後。先區分安全/重大品質紅線、可透過臨時措施控制的風險，以及一般優化項；紅線不能帶病交付。其他問題可評估隔離、返工、分批處理等方案，由製造、品質和工程共同確認風險、責任人與關閉條件。",
                "follow": "客戶願意先收貨，內部品質仍不同意，誰做決定？",
                "follow_answer": "依公司權責與品質放行流程決策；客戶同意不代表內部可以突破安全與品質底線。",
                "risk": "分批處理是可評估方案，不要宣稱公司已採用；PPT沒有不良率或一次驗收率。",
            },
            {
                "q": "製造如何向前端延伸，而不是只增加會議？",
                "line": "提前評審可製造性、物料、調試和售後風險。",
                "answer": "製造向前端延伸，不是多開會，而是在項目早期參與可製造性、裝配、調試、物料與售後風險評審，明確工程、品質、採購和製造的責任與節點。量產和售後問題要再回饋到前端標準。成果應看風險是否更早發現、問題是否減少反覆、方法是否沉澱，而不是會議次數。",
                "follow": "方案已凍結才發現裝配困難，怎麼辦？",
                "follow_answer": "先量化安全、成本與交期影響，提出可選方案並升級決策；同時保留問題記錄，防止下一個專案重複。",
                "risk": "PPT把向前端延伸列為未來規劃，不能說成已有成熟機制。",
            },
        ],
    ),
    (
        "四、人才梯隊與組織文化",
        [
            {
                "q": "AB角如何避免只是名單有備份，實際仍離不開A角？",
                "line": "用能力標準、實作與獨立作業驗證。",
                "answer": "先盤點關鍵崗位和單點風險，為每項工作定義主責、備援與可獨立作業標準。B角要經過帶教、實作、獨立處理三個階段，再用輪值或專案實戰驗證；若仍需要A角全程指導，就不算真正備援。我會先在現有六人調試售後團隊試點，再逐步複製。",
                "follow": "用什麼證據判定B角合格？",
                "follow_answer": "能在A角缺席時按安全、品質與時效要求獨立完成，且問題記錄和主管驗收符合標準。",
                "risk": "AB角屬未來規劃，不要虛報覆蓋率、完成率或已認證人數。",
            },
            {
                "q": "多能工會不會變成每個人都懂一點、但沒有專長？",
                "line": "主技能保深度，備援技能保彈性。",
                "answer": "多能工不是取消專業分工，而是在保留核心專長的前提下增加一到兩項可替補能力。每人要有主技能、備援技能和清楚的合格標準，不追求人人全能。第一批應優先選對交付影響大、目前只有一人會、又能透過實戰培養的技能。",
                "follow": "只能先培養兩個崗位的B角，如何排序？",
                "follow_answer": "按單點風險、缺席影響、任務頻率、培養週期和外部替代難度排序。",
                "risk": "PPT沒有具體技能矩陣與培訓數據，現場要用計畫語氣。",
            },
            {
                "q": "交付壓力大時，如何避免人才培養一直往後排？",
                "line": "把培養嵌入專案，而不是等工作不忙。",
                "answer": "我會把六類重大專案和調試售後工作作為帶教場景：主責示範、B角實作、結案後復盤，再把方法沉澱成標準。每月除了看進度，也要看誰能獨立承擔、誰仍依賴指導。這樣培養與交付是同一件事，不需要另等空檔。",
                "follow": "帶教影響短期效率，你能接受嗎？",
                "follow_answer": "在不突破安全與關鍵交期前提下，接受可控的短期效率成本；先從低風險任務和有緩衝的項目安排實作。",
                "risk": "不要聲稱已有成熟月度人才評審；PPT只提出未來方向。",
            },
            {
                "q": "製造要進度、品質要停下確認，你如何處理衝突？",
                "line": "不靠「多溝通」，靠共同事實和決策原則。",
                "answer": "先把品質風險、交期影響和可行替代方案放在同一個決策框架。涉及安全或重大品質時底線不讓步；一般問題再評估隔離、返工或分段處理。方案要有責任人、完成時點和升級條件。若雙方仍無法接受，就按風險與公司權責向上升級，而不是讓問題停在爭論。",
                "follow": "如果雙方都認為自己是對的，你靠什麼拍板？",
                "follow_answer": "依安全、客戶影響、交期、成本與可逆性排序，並用可驗證資料支持；超出權限時及時升級。",
                "risk": "不要只回答「加強溝通」；必須說清底線、方案和升級條件。",
            },
            {
                "q": "責任文化如何避免變成追責文化？",
                "line": "先看機制，再分辨能力、資源與態度。",
                "answer": "責任文化不是出問題先找人，而是每件事有清楚責任、標準和閉環。像安全庫存，先看清單、分區、盤點和PR機制是否清楚、是否可執行；再判斷是系統缺陷、能力不足、資源不足，還是明知不執行。對事復盤、對人輔導，只有重複且明知違反才進入績效處理。",
                "follow": "什麼情況必須追究個人責任？",
                "follow_answer": "已明確標準與培訓、資源可得，仍故意不執行或隱瞞重大風險，且造成影響時。",
                "risk": "避免空喊「不甩鍋」；要區分系統、能力和態度。",
            },
        ],
    ),
]


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("楊敏銳｜製造部主管晉升課長")
set_repeat_font(hr, size=9, color=MUTED)
hr2 = hp.add_run("    開放式答辯演練手冊")
set_repeat_font(hr2, size=9, color=BLUE, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("內部演練  |  ")
set_repeat_font(fr, size=9, color=MUTED)
add_page_field(fp)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kicker.add_run("MULTI-AGENT DEFENSE SIMULATION")
set_repeat_font(r, size=10, bold=True, color=GOLD)
kicker.paragraph_format.space_after = Pt(14)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("晉升課長開放式答辯問答")
set_repeat_font(r, size=28, bold=True, color=BLUE)
title.paragraph_format.space_after = Pt(8)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("經營主管 × 製造專家 × HR／組織發展 三視角壓力測試")
set_repeat_font(r, size=13, color=DARK_BLUE)
sub.paragraph_format.space_after = Pt(22)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("候選人：楊敏銳  |  目標職位：製造課長  |  版本：2026.07.27")
set_repeat_font(r, size=10.5, color=MUTED)
meta.paragraph_format.space_after = Pt(46)

lead = doc.add_paragraph()
lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = lead.add_run("核心原則：先下結論，再講證據，最後說明下一步；沒有數據時守住邊界，不補造結果。")
set_repeat_font(r, size=11.5, bold=True, color=INK)
shade_paragraph(lead, LIGHT_BLUE)
lead.paragraph_format.left_indent = Inches(0.4)
lead.paragraph_format.right_indent = Inches(0.4)
lead.paragraph_format.space_before = Pt(8)
lead.paragraph_format.space_after = Pt(12)

doc.add_page_break()

h = doc.add_paragraph(style="Heading 1")
r = h.add_run("答辯使用方法")
set_repeat_font(r, size=16, bold=True, color=BLUE)

methods = [
    ("1｜先答結論", "第一句直接回答委員問題，不繞背景。"),
    ("2｜再給證據", "優先使用PPT可守事實：六類專案、六人團隊、24H響應、三階段流程、標準化、安全庫存。"),
    ("3｜補管理思路", "說明如何定義指標、配置資源、設升級條件和形成閉環。"),
    ("4｜守住邊界", "AB角、多能工、人才梯隊、文化建設與向前端延伸屬未來規劃，不能說成既有成果。"),
]
for label, text in methods:
    add_labeled_paragraph(doc, f"{label}｜", text, after=7)

h = doc.add_paragraph(style="Heading 1")
r = h.add_run("高風險口徑")
set_repeat_font(r, size=16, bold=True, color=BLUE)
add_labeled_paragraph(
    doc,
    "不可直接聲稱｜",
    "節省金額、效率提升百分比、交期達成率、良率、客戶滿意度、一次解決率、庫存周轉改善、AB角覆蓋率。",
    fill=LIGHT_GRAY,
)
add_labeled_paragraph(
    doc,
    "可以穩健表達｜",
    "PPT已證明機制與做法；量化效果將透過準時交付、異常關閉、響應與結案、重複發生、庫存和技能矩陣持續驗證。",
    fill=LIGHT_BLUE,
)

qno = 1
for section_index, (section_title, questions) in enumerate(sections):
    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run(section_title)
    set_repeat_font(r, size=16, bold=True, color=BLUE)
    for question in questions:
        if qno == 13:
            doc.add_page_break()
        add_question(doc, qno, question)
        qno += 1

doc.core_properties.title = "楊敏銳晉升課長開放式答辯問答"
doc.core_properties.subject = "多Agent模擬答辯演練手冊"
doc.core_properties.author = "Codex"
doc.core_properties.keywords = "晉升, 課長, 製造, 開放式答辯"
doc.save(OUTPUT)
print(OUTPUT)
