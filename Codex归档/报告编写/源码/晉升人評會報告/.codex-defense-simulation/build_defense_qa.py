from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"C:\Users\lfaf-test\Documents\报告编写\晉升人評會報告\宋佳骥_晉升開放式答辯問答_多主管模拟版.docx")

NAVY = "004B7A"
BLUE = "2FA9D6"
ORANGE = "F28C28"
DARK = "303030"
MID = "666666"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, size=11, bold=False, color=DARK, name="Microsoft YaHei", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID)


def add_labeled_paragraph(doc, label, text, label_color=NAVY, body_color=DARK, before=0, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    lr = p.add_run(label)
    set_run_font(lr, size=11, bold=True, color=label_color)
    br = p.add_run(text)
    set_run_font(br, size=11, color=body_color)
    return p


def add_callout(doc, label, text, fill=PALE, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_question(doc, number, question, answer, followup, response, risk=None):
    p = doc.add_paragraph(style="Question")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Q{number}. {question}")
    set_run_font(r, size=12, bold=True, color=NAVY)
    add_labeled_paragraph(doc, "建议回答：", answer, label_color=NAVY, after=4)
    add_labeled_paragraph(doc, "可能追问：", followup, label_color=GOLD, after=3)
    add_labeled_paragraph(doc, "应对要点：", response, label_color=BLUE, after=5)
    if risk:
        add_labeled_paragraph(doc, "回答风险：", risk, label_color=RED, body_color=RED, after=8)


PANELS = [
    {
        "title": "第一组｜林董视角：为什么晋升、能给组织带来什么",
        "focus": "听的不是你写了多少程序，而是你能否把个人能力放大成组织能力，并对长期结果负责。",
        "questions": [
            ("你做的这些事情，为什么足以支持晋升课长？",
             "我申请晋升，不是因为年资到了，也不只是因为程序写得好，而是工作已经从“自己完成任务”转成“让团队持续完成任务”。技术上，我把PLC和视觉经验做成可复用模块；人员上，我从零建立8人的视觉、AI和机器人团队；现在还要对方案、资源、进度、品质和人才培养负责。晋升后，我要承担的是标准、人才和交付机制的完整责任。",
             "这些工作你现在不晋升也能做，为什么一定要晋升？",
             "不要强调头衔方便；回答正式职责需要与资源调配、骨干培养、跨组协作和结果追责相匹配。"),
            ("Altis和VM看起来都是技术成果，怎么证明你的管理价值？",
             "如果成果只停留在“我会写”，确实只是技术价值。我的管理价值是把一次成功变成团队以后可以重复使用的能力。Altis首线形成架构和功能块，后续复制到另外3条线；VM统一5种相机架设方式后，约50台设备按同一套方法配置。下一步再补齐说明、案例、测试和版本负责人，让成果不依赖我本人。",
             "现在还依赖你，说明管理转化并没有完成吧？",
             "承认只完成了一半：代码复用已经形成，但知识承接尚未完全补齐；下一阶段重点正是完成组织化。"),
            ("从公司长期发展看，你这套工作的价值是什么？",
             "长期价值不是单纯快几天，而是降低公司对少数资深工程师的依赖。把经验变成标准模块、测试方法和版本记录后，新人可以在受控范围内承担更多工作，资深人员把精力放在架构、风险和异常问题上。公司扩大项目时，不必完全按项目数量同比增加资深骨干。",
             "是不是想用新人或AI替代资深工程师？",
             "明确不是替代。安全、架构、验证和异常处理仍需要资深人员；目标是减少重复编写和口头传授。"),
            ("如果你明天离开，这些能力还能留下多少？",
             "坦白说，目前不能说全部留下。程序和部分模块已经可复用，但说明、测试案例和版本责任还不完整。改善方式是每个核心模块明确负责人和备份，程序、说明、案例、测试和版本记录一起交付。真正留下的标准，是负责人不在时，其他人仍能按资料正确使用和验证。",
             "为什么代理课长几年还没有补完？",
             "承认过去更重交付、轻沉淀，不找借口；说明现在已把完成标准从“可运行”升级为“可承接”。"),
            ("你最大的管理短板是什么？短板没补好，为什么现在晋升？",
             "我最大的短板是向上汇报和跨部门推动还不够结构化。过去习惯先解决问题，有时风险和所需支持没有第一时间讲清楚。现在固定用“结论、事实、风险、需要支持、下一步”的结构，并把跨部门事项落到责任人和完成日。晋升不是说明没有短板，而是愿意对短板的改善结果负责。",
             "半年后怎么证明真的改善？",
             "用项目一页报、风险提前暴露记录、责任事项关闭情况和项目复盘作为证据，不临时编目标值。"),
            ("你如何证明自己不是最会救火的人，而是能减少救火的人？",
             "课长的价值不是亲自解决最多故障，而是让问题更早暴露。项目启动时确认接口、边界和验收条件；过程中用里程碑和风险清单；现场问题有责任人、完成日和验证结果；项目后把共性问题更新回模块和案例。以后看的是同类问题是否减少、项目是否更少依赖临时救火。",
             "现在最缺的是什么？",
             "回答核心模块的说明、测试和版本责任，不泛泛说人手不足。"),
        ],
    },
    {
        "title": "第二组｜财务主管视角：数据靠不靠谱、到底值多少钱",
        "focus": "财务最关心数据口径、是否真的省钱、是否只是把人力释放而没有形成现金节省。",
        "questions": [
            ("你的成果到底给公司省了多少钱？",
             "目前能确认的是工时和人力需求变化，不能直接换成金额。我不会拿未确认的工资、加班费或停机成本去乘。可以分三层核算：先算节省的人天，再看是否减少加班、外援或新增人力，最后看是否降低延期和停机损失。财务确认人天单价和成本口径后，我提供项目记录共同核算。",
             "没有金额，怎么证明成果有价值？",
             "补充复用规模、产能释放、交付能力和人才梯队；同时同意后续与财务建立统一核算模板。",
             "不要现场编造金额或ROI。"),
            ("PLC缩短71%、视觉减少75%，比例可靠吗？",
             "两个比例都由前后时间计算：PLC从7天到2天，减少5天，约71%；视觉平均从2天到0.5天，减少1.5天，约75%。这只代表指定开发或配置环节，不代表整个设备项目周期也减少相同比例。完整核算还要保留项目类型、人员和起止时间。",
             "PLC的样本量是多少？",
             "视觉可确认约50台；PLC报告没有统一样本量，不猜数字，后续用项目台账补齐。"),
            ("两个人完成一般需要四到五人的开发，“一般”是谁定的？",
             "Altis是9站整线，实际投入是我和另1名电控；我负责整体架构和主要功能块，首线形成模块，后续复制3条并量产。四到五人是同规模项目的一般配置判断，不应直接当成现金节省。财务效益应以同规模项目的计划工时和实际工时比较，而不是只用人数相减。",
             "有没有工时或考勤记录支持？",
             "如果手边没有，就说需要从项目计划、工时和加班记录核对，不能回答“肯定有”。"),
            ("VM程序约50台，实际释放多少工作量？",
             "按报告平均口径，每台从2天降到0.5天，保守减少1.5人天，约50台估算约75人天。这个只是配置和调试工时估算，不包含后续制程调试，也不等于直接节省75人天工资。更准确要看是否减少加班、缩短交期或支持了更多项目。",
             "人员仍在公司，能叫节省吗？",
             "准确说是产能释放；是否形成现金节省，要看加班、外援、新增人力和项目承接量。"),
            ("AIDC的39台和1,270万件，对财务有什么价值？",
             "39台说明部署规模，1,270万件说明实际使用量，但不能直接等于财务收益。财务价值还要看单台减少多少人工、稼动和维护成本、人员是否真正转岗或减少新增需求。报告能确认单台依项目节省1到2人，但不同项目不能统一按上限相乘。",
             "为什么没有直接给ROI？",
             "各项目人工、稼动和维护口径没有统一；宁可暂不报金额，也不报无法追溯的ROI。"),
            ("AI项目可能要买工具、整理资料、培训人员，怎么控制投资风险？",
             "采用分阶段投入。先整理现有资料和模块，主要投入内部工程时间；再在受控模块上做AI组合和测试；达到可追溯、可重复、测试通过后，才进入整套PLC试点。每阶段看资料完整度、模块复用、生成后修改量和测试通过情况，达不到门槛就暂停扩展。",
             "AI软件预算多少？",
             "报告没有确认金额，不现场猜；先确定试点范围、账号、算力、数据安全和验证人力，再由采购、IIC和财务共同核价。"),
        ],
    },
    {
        "title": "第三组｜现场主管视角：出了问题能不能稳住现场",
        "focus": "现场关心安全、停线、异常定位、交付节奏，以及标准化是否真的能用。",
        "questions": [
            ("Altis中你说自己负责架构和主要功能块，具体做了什么？",
             "我不是只负责某一台设备，而是先把9个站如何配合、信号如何交互、异常如何处理的程序骨架搭起来，再完成主要功能块。另1名电控配合具体开发和调试。我的价值不是说代码全由我写，而是先把规则和模块搭好，让两个人可以分工，后续3条线还能复制。",
             "怎么证明没有把团队成果算到自己身上？",
             "明确说我负责架构、主要功能块和技术判断，具体开发与调试由两人共同完成。"),
            ("PLC开发从7天降到2天，为什么现场调试没有同比下降？",
             "模块化减少的是程序框架、通用动作和接口的重复编写；现场调试仍受机械装配、来料、工艺参数和多专业配合影响，所以不会同比例下降。我不会把开发工时改善夸大成整机交付周期改善。下一步增加模块测试案例、接口检查表和离线验证。",
             "2天的口径是什么？",
             "指同类程序开发工时，不包含完整现场制程调试。"),
            ("VM通用视觉程序换到不同设备上，怎么避免配置错误？",
             "先确认相机安装方式和项目需求是否属于已验证模式，再核对相机参数、坐标方向、标定结果和PLC接口；随后用正常与异常样品测试，最后现场低速验证。若超出通用程序适用边界，就回到方案评审，不能强行套用。",
             "出现问题怎么回退？",
             "保留上一稳定版本、记录变更、验证后再发布，不在量产版本上直接试改。"),
            ("设备现场突然出现异常，你作为课长怎么处理？",
             "先保安全、保产品，必要时停止自动运行；再判断是程序、机械、视觉、网络还是制程问题；保留报警、日志、程序版本和现场条件；然后明确临时恢复和正式改善方案。恢复后还要复盘原因、责任人、完成时间和验证结果。",
             "生产很急，能不能先取消互锁？",
             "不能绕过安全互锁。临时措施也必须有主管确认、替代防护、记录和恢复计划。"),
            ("一套方案怎么复制到其他厂，而不是把问题也复制过去？",
             "复制前分清通用部分和厂区、设备、制程特有部分。通用模块可复用，但IO、网络、权限、安全条件、工艺参数和验收标准必须重新确认。上线问题要由模块负责人判断是个案还是共性，只有验证后的共性改善才能回写标准版本。",
             "谁决定回写标准版本？",
             "核心模块负责人评估影响，经过测试验证后发布；现场不能直接改标准版本。"),
            ("你有没有失败案例？从中学到什么？",
             "我不会编一个没有确认的重大事故。能够明确承认的不足是：模块代码做到了可复用，但说明、测试案例和版本责任没有同步补齐，形成熟悉的人会用、新人仍要问骨干。这说明技术完成不等于组织承接完成，今后五项资料齐全才算真正完成。",
             "为什么没有早点做？",
             "承认过去优先解决交付和人力，管理取舍不完整；现在把文件化纳入交付标准。"),
        ],
    },
    {
        "title": "第四组｜IIC主管视角：网络、权限、AI与信息安全",
        "focus": "IIC会关注谁有权限、数据是否外泄、远程操作如何追溯，以及AI能否受控。",
        "questions": [
            ("远程支援时，电控和IIC的权限边界在哪里？",
             "电控负责设备侧联网需求、IP和接口规格，以及调试需要观察或操作的内容；账号、访问权限和安全策略由IIC统一管理，电控不能为了方便自行绕过。远程前确认设备状态、操作窗口和现场配合人员，涉及高权限或生产风险时由IIC与现场共同确认。",
             "为什么不能由电控自己开权限？",
             "网络权限、安全审计和跨厂访问属于IIC职责，技术上能连上不等于可以随意操作。",
             "此案例属于答辩备用，不在正式报告中主动展开。"),
            ("远程调试发生误操作，责任怎么划分？",
             "先通过流程减少争议：明确申请人、操作人、现场确认人、时间窗口和操作范围；变更前保存版本，操作中留日志，完成后现场验证。未经授权或超范围操作由操作人员负责；需求、批准或现场状态确认不清，也要追溯流程责任。",
             "紧急停线还要不要走流程？",
             "可以启用紧急流程，但至少要有现场确认、许可边界、操作记录和事后复盘。"),
            ("AI可调用模块库和普通程序文件夹有什么区别？",
             "普通文件夹只是把程序放在一起；AI可调用模块库要求每个模块的功能、输入输出、适用条件、限制、测试案例和版本都描述清楚。AI只能从经过验证的模块中选择和组合，工程师再负责导入、编译、功能测试、安全验证和现场验收。",
             "谁判断AI选错了模块？",
             "系统设计人员和模块负责人审核，测试验证人员用案例和异常工况确认。"),
            ("AI生成PLC程序，怎样保证设备安全？",
             "当前AI只做到可导入、可编译，不代表功能正确或可以量产。以后限制AI只能调用经过验证、版本受控的模块，再做编译检查、离线测试、IO与互锁检查、异常工况测试和现场低速验证。安全回路、危险动作和关键互锁必须人工逐项复核并签字。",
             "如果AI提高不了效率怎么办？",
             "先用于资料整理和模块调用，小范围试点；是否扩大以验证结果为准，不为AI牺牲安全。"),
            ("使用AI整理程序资料，会不会把公司代码或设备数据泄露？",
             "必须先由IIC确认允许使用的工具、账号、数据范围和存储位置。未经批准的外部模型不能上传公司源代码、客户资料、网络配置或设备数据。试点应从脱敏资料和非敏感模块开始，保留访问和版本记录。技术试点必须服从公司信息安全规则。",
             "如果现有工具无法满足保密要求怎么办？",
             "停止上传敏感资料，改用公司批准的环境或只做本地、脱敏验证，不能绕过安全要求。"),
        ],
    },
    {
        "title": "第五组｜人事主管视角：你会不会带人、管人、承担责任",
        "focus": "人事关注从工程师到主管的角色变化、绩效、公平、授权、冲突和继任。",
        "questions": [
            ("从工程师变成课长，最大的变化是什么？",
             "工程师主要对自己负责的程序和设备结果负责；课长要对整个团队的交付负责。重点从“我能不能解决”转为“方案有没有提前评审、人员怎么安排、风险有没有暴露、团队离开我还能不能完成”。重大问题我仍会介入，但不能每次都靠我救火。",
             "现场重大问题你还会不会亲自上手？",
             "可以介入，但同步安排成员参与，处理后沉淀成模块、案例或检查表。"),
            ("你过去偏向自己解决问题，凭什么能带好19人团队？",
             "这个习惯确实需要改变，但我已经不只是个人开发。我参与招聘并从零建立8人的视觉、AI和机器人团队，通过入门带教、自主研究和每周复盘，目前4人能开发、另外4人侧重应用调试。晋升后要把培养、授权和复盘进一步制度化。",
             "为什么另外4人还不能开发？",
             "多数人员从应届生培养，成长需要过程；应用调试也是交付能力，下一步建立从应用、单模块、单机到整站的升级路径。"),
            ("你会怎样管理下属绩效？",
             "不会只看加班或写了多少程序，而看结果和成长。结果包括节点、质量和风险报告；成长包括能否独立承担更复杂任务、完善文档和带新人。项目开始先讲清目标和验收条件，过程中按节点检查，结束后复盘。表现不达标时先区分能力、任务还是态度问题。",
             "辅导后仍没有改善怎么办？",
             "设明确改善期限和可验证目标；仍未改善就调整任务或岗位，态度和纪律问题按制度处理。"),
            ("两名骨干发生冲突，你怎么处理？",
             "先把争论拉回项目目标，分别听清依据，判断是技术、资源还是责任边界。技术问题用数据和测试判断，资源按优先级决定，责任重新写清。最终由我做决定并说明理由，决定后必须执行，项目结束再复盘机制问题。",
             "其中一人仍不接受怎么办？",
             "可以保留意见，但不能影响交付；要求用事实提出异议，持续影响合作则按行为和绩效处理。"),
            ("你怎么授权，才能不包办也不放任？",
             "授权要讲清目标、边界、权限和检查点。新人从小范围任务开始，骨干可负责完整模块或项目；我在方案、关键节点和高风险动作上检查。一般问题先让负责人提出方案，涉及安全、客户节点或重大资源冲突时我再介入决策。",
             "下属做错导致项目受影响，责任算谁的？",
             "主管对授权和检查机制承担管理责任，下属对职责内执行负责，重点是复盘边界和检查点。"),
            ("如果两名关键骨干同时离职，你怎么办？",
             "目前新人对少数骨干仍有依赖，这是报告已经承认的风险。降低影响要靠三件事：关键模块有说明、案例、测试和版本记录；每个模块有主负责人和备份；通过轮岗和带教让知识不只在一个人手里。人员离开不能等于技术和项目一起中断。",
             "骨干不愿意分享怎么办？",
             "先了解原因；文档、模块维护和带教纳入职责评价，同时给分享者更大的技术责任和发展机会。"),
        ],
    },
    {
        "title": "第六组｜模具主管视角：跨专业接口怎么提前说清楚",
        "focus": "模具主管会从定位、公差、变更、接口扯皮和联合验收来判断你的协同能力。",
        "questions": [
            ("设备开发前，需要模具部门提供哪些信息？",
             "至少要明确产品定位方式、基准面、允许公差、取放空间、治具开合方式和换型范围。电控与视觉还要知道到位信号、异常状态和人工处理方式。我的做法是在方案阶段把机械、模具、电控和视觉边界列成清单，共同确认节拍、信号和验收条件。",
             "前期产品资料还不完整怎么办？",
             "区分已确认和暂定条件，对暂定项设冻结时间和变更责任，高风险接口先验证。"),
            ("后期模具变更导致PLC或视觉重做，你怎么处理？",
             "先判断变更影响定位、动作顺序、信号接口、视觉参数还是安全互锁，再评估工时和节点影响。能通过参数或标准模块调整的控制在模块内，涉及架构变化则重新评审。变更原因、影响范围、责任人和完成时间必须同步各专业。",
             "延期责任算谁的？",
             "按变更原因和已确认的冻结条件判断；对外先共同解决交付，对内依据记录复盘责任。"),
            ("模具说是设备动作问题，电控说是模具精度问题，怎么处理？",
             "不能靠争论，要把设备动作、模具状态和产品结果拆开验证。先确认定位与基准稳定，再查传感器、动作顺序和参数，通过重复和对照测试看问题跟哪个条件变化。我的责任是把争论变成可验证项目。",
             "双方都坚持不是自己的问题怎么办？",
             "统一测试条件和判定标准，指定项目负责人汇总数据，必要时请质量或工艺共同判定。"),
            ("怎么保证模具、设备、电控和视觉接口不会到现场才暴露？",
             "分三阶段检查：方案阶段确认定位、动作、信号、节拍和异常处理；开发阶段按接口清单确认各专业输入输出；现场前针对正常、异常、换型和安全场景联合验证。无法提前验证的要明确风险、现场负责人和预案。",
             "哪个部门负责最终接口验收？",
             "各专业对本专业输出负责，项目负责人拉通整机接口和最终交付，验收必须有共同条件和记录。"),
            ("标准模块会不会限制模具或制程创新？",
             "标准模块是减少重复工作，不是强迫所有项目完全一样。稳定、共性的功能使用标准模块；遇到新制程或特殊机构，可以建立项目专用方案，验证成熟后再评估是否纳入标准。标准必须有适用边界，不能为了复用牺牲制程效果。",
             "谁决定特殊方案是否进入标准库？",
             "由系统设计、核心模块负责人和相关专业共同评审，经过项目验证后发布。"),
        ],
    },
]


QUICK_FIRE = [
    ("为什么晋升？", "因为我已经不只解决个人技术问题，而是在建立标准、培养人才、配置资源并对团队交付负责。"),
    ("最大优势？", "懂现场和技术，也有把技术变成标准、人才和交付结果的实践。"),
    ("最大短板？", "向上汇报、跨部门推动，以及模块说明和测试承接还不完整。"),
    ("最硬的数字？", "Altis两人完成通常需四到五人的开发；PLC开发由7天降到2天。"),
    ("最值得骄傲的成果？", "从零建立8人的视觉、AI和机器人团队，并让团队支撑跨厂设备应用。"),
    ("核心模块负责人是什么？", "不是所有程序都由他写，而是长期对某类模块的质量、说明、测试和版本负责。"),
    ("AI现在做到什么？", "目前只做到可导入、可编译，不等于可运行、可量产，必须人工验证。"),
    ("AI会取代工程师吗？", "不会简单替代，工程师工作会更多转向需求、架构、测试、安全和异常处理。"),
    ("数据不知道怎么办？", "明确说没有准确数据、不做猜测，再讲能确认的口径和后续补齐方法。"),
    ("如果没有晋升？", "继续按课长标准做事，确认差距，用下一阶段可验证结果证明改善。"),
]


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 18, 10),
    ("Heading 2", 13, NAVY, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    st = styles[style_name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Question" not in styles:
    question_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
else:
    question_style = styles["Question"]
question_style.font.name = "Microsoft YaHei"
question_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
question_style.font.size = Pt(12)
question_style.font.bold = True
question_style.font.color.rgb = RGBColor.from_string(NAVY)
question_style.paragraph_format.space_before = Pt(10)
question_style.paragraph_format.space_after = Pt(5)
question_style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("晋升开放式答辩｜多主管模拟训练")
set_run_font(hr, size=9, bold=True, color=MID)
footer = section.footer
fp = footer.paragraphs[0]
fr = fp.add_run("宋佳骥｜工程师 → 课长    ")
set_run_font(fr, size=9, color=MID)
add_page_number(fp)

# Editorial-cover opening.
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("晋升开放式答辩")
set_run_font(r, size=30, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("林董 · 现场 · 财务 · 人事 · IIC · 模具主管")
set_run_font(r, size=16, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("多智能体模拟问答与压力追问训练")
set_run_font(r, size=13, color=MID)
p.paragraph_format.space_after = Pt(60)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("依据：宋佳骥晋升人评会报告R5、既有开放式答辩问答及用户确认数据")
set_run_font(r, size=10, color=MID)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("版本：2026年7月｜答案建议控制在30–60秒")
set_run_font(r, size=10, color=MID)
doc.add_page_break()

doc.add_heading("使用说明｜不要逐字背，记住回答骨架", level=1)
add_callout(doc, "答题骨架：", "先讲结论（10秒）→ 给事实、数字或案例（20–30秒）→ 承认边界并说下一步（10–20秒）。")
add_callout(doc, "关键原则：", "不知道的数据不猜；被质疑先承认口径；成果归属讲清“我负责什么、团队完成什么”；最后落到组织能力。", fill="FFF8E8", accent=GOLD)
add_callout(doc, "正式报告边界：", "经营管理能力页未放入正式报告。K7调度、远程支援等案例只作为被追问时的备用证据，不主动抢先展开。", fill="FDECEC", accent=RED)

doc.add_heading("答辩数字速记", level=2)
facts = [
    ("经历", "2011年毕业后从事现场设备维修；2015年加入LFAF；2021年起代理课长"),
    ("团队", "含本人19人：传统PLC 10人；视觉／AI／机器人8人"),
    ("Altis", "9站整线；2人完成一般需4–5人的开发；复制3条，共4条量产线"),
    ("PLC模块化", "同类开发由7天降到2天，缩短约71%"),
    ("VM视觉", "5种相机架设模式；约50台设备；平均2天降到0.5天；相同架设约3小时"),
    ("AIDC", "跨厂部署39台；累计实际检测1,270万+；单台依项目节省1–2人"),
    ("人才", "8人团队中4人具备开发能力，4人侧重应用调试"),
    ("AI边界", "当前仅可导入、可编译；量产前必须完成功能、异常与安全验证"),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [1700, 7660])
for i, text in enumerate(("主题", "确认口径")):
    set_cell_shading(table.rows[0].cells[i], LIGHT)
    p = table.rows[0].cells[i].paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
for label, value in facts:
    cells = table.add_row().cells
    set_table_geometry(table, [1700, 7660])
    for cell, text, bold in ((cells[0], label, True), (cells[1], value, False)):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=10.5, bold=bold, color=DARK)

question_number = 1
for panel in PANELS:
    doc.add_page_break()
    doc.add_heading(panel["title"], level=1)
    add_callout(doc, "评委关注：", panel["focus"], fill=PALE, accent=NAVY)
    for question, answer, followup, response, *risk in panel["questions"]:
        add_question(doc, question_number, question, answer, followup, response, risk[0] if risk else None)
        question_number += 1

doc.add_page_break()
doc.add_heading("压力快问快答｜每题20秒", level=1)
for idx, (question, answer) in enumerate(QUICK_FIRE, 1):
    p = doc.add_paragraph(style="Question")
    r = p.add_run(f"{idx}. {question}")
    set_run_font(r, size=12, bold=True, color=NAVY)
    add_labeled_paragraph(doc, "短答：", answer, label_color=BLUE, after=8)

doc.add_heading("财务数据安全句式", level=2)
for label, text in [
    ("金额不确定：", "这个金额目前没有经过财务确认，我不做猜测；我能确认的是工时、数量和应用范围。"),
    ("比例被质疑：", "这个比例只对应具体开发或配置环节，不代表整个项目周期，我先把口径讲清楚。"),
    ("样本量不足：", "目前报告没有统一样本数，我不现场补数字；后续用项目台账和工时记录补齐。"),
    ("成果归属：", "具体开发由团队共同完成，我负责方向、架构或资源、关键判断和交付结果。"),
]:
    add_labeled_paragraph(doc, label, text, label_color=GOLD, after=6)

doc.add_heading("不要这样说", level=2)
avoid_rows = [
    ("“这个我不清楚。”", "“这个指标目前没有准确数据，我不做猜测；我能确认的是……”"),
    ("“都是下面的人做的。”", "“我负责方向、资源和结果，具体开发由团队成员承担。”"),
    ("“AI以后可以全部自动写。”", "“当前仅可导入编译，先从受控模块调用和人工验证开始。”"),
    ("“我们没有问题。”", "“当前主要风险是技术承接薄弱，我已明确补齐机制。”"),
    ("“因为人不够，只能这样。”", "“资源不足是事实，我的责任是提前预警并给出调度方案。”"),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [3900, 5460])
for i, text in enumerate(("避免说法", "建议说法")):
    set_cell_shading(table.rows[0].cells[i], LIGHT)
    r = table.rows[0].cells[i].paragraphs[0].add_run(text)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
for bad, good in avoid_rows:
    cells = table.add_row().cells
    set_table_geometry(table, [3900, 5460])
    for cell, text in zip(cells, (bad, good)):
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, color=DARK)

doc.add_heading("15分钟模拟顺序", level=2)
for step in [
    "第1轮（3分钟）：林董问“为什么晋升”和“长期价值”。",
    "第2轮（4分钟）：财务追问数据口径、金额和AI投入。",
    "第3轮（4分钟）：现场与IIC追问异常、安全、权限和AI边界。",
    "第4轮（3分钟）：人事与模具追问带人、冲突、变更和跨部门接口。",
    "最后1分钟：随机抽取3道快问快答；每题不超过20秒。",
]:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(step)
    set_run_font(r, size=11, color=DARK)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
