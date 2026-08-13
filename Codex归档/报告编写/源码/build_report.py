from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\lfaf-test\Documents\报告编写")
OUT = ROOT / "报告" / "自动化设备研发与交付流程梳理报告.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "25364A"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            set_cell_margins(cell)


def style_cell_text(cell, size=9.3, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        for r in p.runs:
            set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        set_cell_shading(c, PALE_BLUE)
        style_cell_text(c, size=9.2, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=font_size)
    set_table_geometry(table, widths)
    set_table_borders(table)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0,
             align=WD_ALIGN_PARAGRAPH.LEFT, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    c = table.cell(0, 0)
    set_cell_shading(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label + "  ")
    set_run_font(r, 10.2, True, DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, 10.2, False, INK)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=1)
    add_para(doc, "", after=1)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_page_break(doc):
    # Allow Word to paginate naturally; headings already keep with following content.
    return None


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.95)
sec.right_margin = Inches(0.95)
sec.header_distance = Inches(0.36)
sec.footer_distance = Inches(0.36)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for level, size, before, after, color in [
    (1, 16, 14, 7, BLUE), (2, 13, 10, 5, BLUE), (3, 11.5, 7, 3, DARK_BLUE)
]:
    st = styles[f"Heading {level}"]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Running header/footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("内部流程梳理｜依据手写原稿整理")
set_run_font(hr, 8.5, False, MUTED)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("自动化设备研发与交付流程梳理报告")
set_run_font(fr, 8.3, False, MUTED)

# Masthead
add_para(doc, "流程管理报告", size=10, bold=True, color=BLUE, after=4)
add_para(doc, "自动化设备研发与交付流程梳理报告", size=23, bold=True, color=INK, after=5)
add_para(doc, "覆盖需求确认、BOM、制图、程序开发、调试交付与售后闭环", size=12.5, color=MUTED, after=14)

meta = doc.add_table(rows=3, cols=2)
meta_data = [
    ("资料来源", "现场手写流程草图"),
    ("整理日期", "2026年7月20日"),
    ("文档状态", "初版｜待业务负责人复核"),
]
for i, (k, v) in enumerate(meta_data):
    meta.cell(i, 0).text = k
    meta.cell(i, 1).text = v
    set_cell_shading(meta.cell(i, 0), LIGHT)
    style_cell_text(meta.cell(i, 0), bold=True, color=DARK_BLUE)
    style_cell_text(meta.cell(i, 1))
set_table_geometry(meta, [1800, 7560])
set_table_borders(meta, color="D9DEE7", size=5)
add_para(doc, "", after=2)

add_callout(doc, "核心结论", "原稿描述的并非单一部门动作，而是一条以客户需求为起点、以交付和维修复盘为终点的跨专业闭环。当前重点应放在统一阶段门、明确输出物、建立BOM与图纸版本基线，并让现场问题能够回写到设计和标准库。")

add_heading(doc, "一、背景与目的", 1)
add_para(doc, "本报告将手写草图中的零散信息整理为可执行的自动化设备研发与交付流程。整理范围包括前期需求与方案、立项开发、BOM建立、电气与流程制图、程序设计、调试入库、出货售后及维修总结。报告不对草图未给出的组织名称、审批权限、周期和量化指标作推定。")

add_heading(doc, "二、端到端流程总览", 1)
workflow_rows = [
    ("1", "需求确认", "澄清客户需求、边界、场景与验收口径", "需求确认记录"),
    ("2", "方案沟通与评审", "事业部/相关专业共同评估方案可行性", "评审意见、方案基线"),
    ("3", "立项开发", "形成开发任务，组织机械、电气、软件协同", "立项信息、任务分解"),
    ("4", "BOM与制图", "建立物料清单，输出电气图、流程图等", "受控BOM、设计图纸"),
    ("5", "程序设计", "完成PLC、CCD、机器人、HMI程序", "程序包、参数与版本"),
    ("6", "调试与入库", "联调、问题整改、验证并办理入库", "调试记录、入库资料"),
    ("7", "出货与售后", "交付、安装/支持、现场问题响应", "交付清单、售后记录"),
    ("8", "维修总结", "汇总故障、分析原因并沉淀改进", "维修报告、改进项"),
]
add_table(doc, ["阶段", "名称", "关键活动", "建议输出物"], workflow_rows, [720, 1640, 4180, 2820], 8.9)

add_heading(doc, "三、流程控制原则", 1)
principles = [
    ("单一基线", "需求、BOM、图纸和程序均应有唯一有效版本，变更后同步更新关联资料。"),
    ("阶段门控制", "上阶段关键输出未确认时，不宜直接进入采购、加工或批量程序开发。"),
    ("跨专业一致", "机械、电气、软件和现场调试对设备配置、点位、接口和安全逻辑采用同一口径。"),
    ("闭环复用", "售后与维修问题要回写到BOM、图纸、程序模板及标准件库。"),
]
add_table(doc, ["原则", "管理要求"], principles, [1900, 7460], 9.4)

add_page_break(doc)
add_heading(doc, "四、BOM建立与物料管理", 1)
add_para(doc, "BOM是采购、装配、调试和售后追溯的共同数据基础。原稿将物料来源大致区分为通用/常规件、已确定件、非标件和新增选型件，并提示可结合ERP或类似系统录入。建议把这种分类固化为统一字段。")

bom_rows = [
    ("通用/常规件", "常用标准件、基础件等", "优先引用标准物料编码与合格供应来源；减少重复建码。"),
    ("已确定件", "已有规格、型号或客户指定件", "确认规格、品牌、数量、交期及替代限制。"),
    ("非标件", "需加工或定制的结构件/安装件", "关联图号、材料、表面处理、版本和加工要求。"),
    ("新增选型件", "首次采用或需技术比选的器件", "记录选型依据、接口条件、技术参数、替代方案。"),
]
add_table(doc, ["物料类型", "典型内容", "控制要点"], bom_rows, [1800, 2840, 4720], 9.0)

add_heading(doc, "BOM建议字段", 2)
add_para(doc, "至少包括：层级、物料编码、名称、规格型号、品牌/供应商、数量、单位、图号或资料链接、所属模块、关键参数、替代关系、采购/自制属性、版本、状态及备注。对于安全件、长交期件和关键接口件，应增加显著标识。")

add_heading(doc, "BOM建立步骤", 2)
bom_steps = [
    ("01", "模块拆分", "按设备功能、工位或专业拆分BOM层级。"),
    ("02", "物料匹配", "先查标准库与历史项目，再决定引用、替代或新增。"),
    ("03", "新增选型", "核实技术参数、接口、供货周期和现场适用性。"),
    ("04", "专业校核", "机械、电气、软件共同检查数量、接口与配置一致性。"),
    ("05", "版本发布", "完成审批后形成受控基线，并同步采购与装配。"),
    ("06", "变更闭环", "任何替换或数量变化均记录原因、影响和批准信息。"),
]
add_table(doc, ["序号", "步骤", "要求"], bom_steps, [760, 1700, 6900], 9.1)

add_heading(doc, "可复用资料来源", 2)
add_para(doc, "草图提及官网、内部资料、网络方案、电商链接及既往量产/项目经验。建议规定使用优先级：企业标准库与历史验证件优先，其次为厂商官网与正式技术资料；网络方案和电商页面只能作为线索，关键参数必须回到可追溯的正式资料确认。")

add_page_break(doc)
add_heading(doc, "五、设计制图与文件输出", 1)
add_heading(doc, "1. 电气类图纸", 2)
add_para(doc, "根据草图，电气设计输出应覆盖电路/接线关系、元件排布、线路连接、接线端子、设备配置及系统架构。图纸应与BOM、I/O点表和程序变量命名保持一致。")

drawing_rows = [
    ("电气原理/接线图", "电源、控制回路、安全回路、信号连接", "回路完整、线号统一、器件代号与BOM一致"),
    ("元件排布图", "柜内、操作台或设备端元件位置", "安装空间、散热、维护与走线可达性"),
    ("端子与线缆资料", "端子排、线号、线缆规格及去向", "现场可按表接线并快速排障"),
    ("设备配置/架构图", "PLC、CCD、机器人、HMI等网络与接口", "通信协议、地址、接口边界清晰"),
]
add_table(doc, ["图纸/资料", "主要内容", "检查重点"], drawing_rows, [2200, 3440, 3720], 9.0)

add_heading(doc, "2. 流程与操作类图纸", 2)
add_para(doc, "流程图和操作步骤图用于表达设备动作顺序、条件分支、异常处理与人机交互，是程序设计思路的直接输入。建议至少覆盖正常启动、自动循环、暂停/复位、故障恢复和安全停机等场景。")

add_heading(doc, "3. 文件发布与制造接口", 2)
add_para(doc, "设计文件发布时应统一输出受控源文件和PDF版本。提供制造或接线使用的PDF时，应同时带版本号、发布日期、适用设备/项目和变更说明，避免车间或外协单位使用过期图纸。草图中还提到线路标识制作，建议将线号、端子号和标签清单纳入同一次发布。")

add_callout(doc, "发布检查", "BOM、图纸、端子/线号清单、I/O点表、程序变量表中的名称和编号必须能够互相对应；任何一处变更都要评估对其他文件的影响。", fill=PALE_GOLD)

add_heading(doc, "六、程序设计与联调", 1)
program_rows = [
    ("PLC/PC", "顺序控制、I/O、互锁、安全条件、设备通信", "点位表、报警码、状态机、异常恢复"),
    ("CCD/视觉", "图像采集、定位/检测、结果判定与数据交互", "光源与镜头、配方、阈值、误检漏检验证"),
    ("Robot", "轨迹、坐标系、动作节拍、抓取/放置及互锁", "工具/工件坐标、禁入区、碰撞与恢复策略"),
    ("HMI", "操作画面、参数、报警、权限与维护入口", "状态可视、操作防错、报警可定位、参数可追溯"),
]
add_table(doc, ["系统", "主要任务", "关键关注点"], program_rows, [1500, 4000, 3860], 8.9)

add_heading(doc, "联调建议顺序", 2)
add_para(doc, "建议按照“离线检查—单元测试—接口联调—整机空载—带料验证—异常与安全测试—验收”的顺序推进。每次测试均记录前提、步骤、结果、问题、责任接口和复测结论。程序发布需保留源文件、编译/运行版本、参数备份及回滚版本。")

add_page_break(doc)
add_heading(doc, "七、调试、入库、交付与售后", 1)
delivery_rows = [
    ("调试准备", "图纸/BOM/程序版本一致；物料齐套；安全条件具备", "调试准备检查表"),
    ("功能验证", "正常流程、节拍、精度、报警、互锁、断电恢复", "测试记录与问题清单"),
    ("问题整改", "定位机械、电气、软件或工艺原因并复测", "整改闭环记录"),
    ("入库/放行", "资料齐套、关键指标达标、遗留项有明确处置", "入库/放行记录"),
    ("出货交付", "设备、备件、资料、程序备份和培训内容同步移交", "交付清单、签收记录"),
    ("售后支持", "安装调试、故障响应、远程/现场支持", "售后工单、服务记录"),
    ("维修总结", "统计故障现象、原因、措施、复发风险和标准化建议", "维修总结、改进任务"),
]
add_table(doc, ["环节", "控制重点", "建议记录"], delivery_rows, [1800, 4760, 2800], 8.9)

add_heading(doc, "八、职责分工建议", 1)
role_rows = [
    ("项目/业务负责人", "组织需求确认、方案评审、节点协调与客户接口"),
    ("机械设计", "结构方案、非标件图纸、装配可行性与机械风险"),
    ("电气设计", "电气BOM、原理/接线/排布图、I/O及系统架构"),
    ("软件/自动化", "PLC、CCD、机器人、HMI程序及版本管理"),
    ("装配与调试", "按受控资料实施装配、联调、记录问题并复测"),
    ("采购/供应链", "按受控BOM采购，反馈交期、替代与质量风险"),
    ("售后/维修", "记录现场问题，完成处理并推动经验回写"),
]
add_table(doc, ["角色", "主要职责"], role_rows, [2500, 6860], 9.2)

add_heading(doc, "九、建议建立的受控资料包", 1)
package_rows = [
    ("需求与方案", "需求确认单、技术协议、方案及评审记录"),
    ("物料", "分层BOM、选型资料、替代清单、关键件清单"),
    ("图纸", "机械图、电气图、端子/线号清单、架构图、流程图"),
    ("软件", "PLC/PC、CCD、Robot、HMI源文件与发布包、参数备份"),
    ("验证", "调试记录、测试报告、问题清单、复测与验收记录"),
    ("交付与服务", "交付清单、培训资料、售后工单、维修总结"),
]
add_table(doc, ["资料类别", "主要内容"], package_rows, [2500, 6860], 9.2)

add_page_break(doc)
add_heading(doc, "十、主要风险与改进建议", 1)
risk_rows = [
    ("需求口径不一", "开发返工、验收争议", "需求确认后形成基线；变更需评估影响并留痕。"),
    ("BOM与图纸不同步", "错采、错装、现场返工", "设置跨专业校核与发布前一致性检查。"),
    ("程序缺少版本控制", "难以复现问题或快速回退", "统一命名、备份、发布记录和回滚机制。"),
    ("网络资料直接用于选型", "参数错误、供货或质量不可控", "以厂商正式资料和内部验证结果为准。"),
    ("售后问题未回写", "同类故障重复发生", "维修总结形成标准库更新或设计变更任务。"),
]
add_table(doc, ["风险", "可能影响", "改进建议"], risk_rows, [2100, 2800, 4460], 8.9)

add_heading(doc, "十一、近期落地行动", 1)
action_rows = [
    ("1", "统一八阶段流程及阶段门定义", "形成一页流程图和发布规则", "待指定"),
    ("2", "确定BOM字段、分类与编码规则", "发布BOM模板及填写说明", "待指定"),
    ("3", "建立图纸与程序命名/版本规范", "实现资料可追溯、可回滚", "待指定"),
    ("4", "制作调试、交付和维修记录模板", "统一问题闭环证据", "待指定"),
    ("5", "选取一个在制项目试运行", "验证流程并根据反馈修订", "待指定"),
]
add_table(doc, ["序号", "行动", "预期结果", "责任人"], action_rows, [700, 3300, 3960, 1400], 8.9)

add_heading(doc, "十二、待确认事项", 1)
add_para(doc, "以下信息在手写原稿中存在歧义或未给出，正式发布前建议由业务负责人确认：")
confirm_rows = [
    ("流程名称", "前期阶段中“事业部沟通/方案评审”等具体名称与顺序。"),
    ("系统名称", "BOM录入所用系统是否为ERP、PDM或其他内部平台。"),
    ("审批权限", "需求、方案、BOM、图纸、程序、入库和出货的批准角色。"),
    ("交付标准", "节拍、精度、良率、安全、资料齐套等具体验收指标。"),
    ("版本规则", "项目号、图号、程序号、版本号及变更单的编码方式。"),
    ("时限与责任人", "各阶段目标周期、责任部门和异常升级路径。"),
]
add_table(doc, ["确认主题", "需确认内容"], confirm_rows, [2200, 7160], 8.8)

# Keep headings with following content and avoid orphan table rows where possible.
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.widow_control = True

doc.core_properties.title = "自动化设备研发与交付流程梳理报告"
doc.core_properties.subject = "依据手写流程草图整理的流程管理初稿"
doc.core_properties.author = ""
doc.core_properties.keywords = "自动化设备,研发流程,BOM,制图,程序设计,调试,交付,售后"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
